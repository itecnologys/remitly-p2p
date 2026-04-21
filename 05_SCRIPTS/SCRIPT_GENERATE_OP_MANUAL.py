import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_markdown_to_doc(md_path, doc):
    if not os.path.exists(md_path):
        print(f"Erro: Arquivo não encontrado - {md_path}")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        
        # Handle Alerts and horizontal rules
        if line.startswith('---'):
            continue
        if line.startswith('> [!'):
            line = line.replace('> [!NOTE]', 'OBSERVAÇÃO:').replace('> [!IMPORTANT]', 'IMPORTANTE:').replace('> [!WARNING]', 'AVISO:')

        if not line:
            doc.add_paragraph()
            continue

        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('|'):
            # Basic table handling (simplified)
            p = doc.add_paragraph(line)
            p.style.font.name = 'Courier New'
            p.style.font.size = Pt(9)
        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def generate_handbook():
    print("Iniciando geração do Manual Operacional FLUXUS...")
    doc = Document()
    
    # Text Styles
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    repo_root = Path(__file__).resolve().parents[1]
    # Content files in order (nomes canónicos no repo actual)
    files = [
        repo_root / "04_PROCESS" / "PROCESS_TECHNICAL_GLOSSARY.md",
        repo_root / "03_TECHNICAL" / "TECH_MATCHING_ENGINE_PROTOCOL.md",
        repo_root / "04_PROCESS" / "PROCESS_LP_CASH_CYCLE.md",
        repo_root / "04_PROCESS" / "PROCESS_LP_DASHBOARD_SPEC.md",
    ]

    for i, file_path in enumerate(files):
        add_markdown_to_doc(str(file_path), doc)
        if i < len(files) - 1:
            doc.add_page_break()

    output_path = repo_root / "04_PROCESS" / "PROCESS_OPERATIONAL_HANDBOOK_AUTOGEN.docx"
    doc.save(str(output_path))
    print(f"Manual salvo com sucesso em: {output_path.resolve()}")

if __name__ == "__main__":
    generate_handbook()
