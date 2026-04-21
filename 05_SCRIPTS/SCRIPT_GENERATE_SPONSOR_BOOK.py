import os
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn

def set_table_header_background(cell, color_hex):
    """Sets the background color of a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_hatched_shading(cell, color_hex):
    """Adds a subtle shading/hatch to a cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'pct10') # Subtle hatch pattern
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')
    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def generate_book(md_path, docx_path, cover_path):
    print("Iniciando geração do Livro Estratégico...")
    doc = Document()

    # Define Global Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(11)

    # Cover Page
    if os.path.exists(cover_path):
        doc.add_picture(cover_path, width=Inches(6.0))
        doc.add_page_break()

    # Introduction / Header
    section = doc.sections[0]
    header = section.header
    ht = header.paragraphs[0]
    ht.text = "FLUXUS STRATEGIC ANNALS - CONFIDENCIAL"
    ht.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer
    ft = footer.paragraphs[0]
    ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ft.add_run()
    run.add_text("Página ")
    add_page_number(run)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    table_data = []
    in_table = False

    for line in lines:
        line = line.strip()
        
        # Handle Headings
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:], level=1)
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=2)
            continue
            
        # Handle Tables (Simple Markdown Table parser)
        if line.startswith('|'):
            if '---' in line:
                continue
            cells = [c.strip() for c in line.split('|') if c.strip()]
            table_data.append(cells)
            in_table = True
            continue
        elif in_table and not line:
            # Create the Word Table
            if table_data:
                rows = len(table_data)
                cols = len(table_data[0])
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                
                for r_idx, row_data in enumerate(table_data):
                    row = table.rows[r_idx]
                    for c_idx, text in enumerate(row_data):
                        cell = row.cells[c_idx]
                        cell.text = text
                        # Styling
                        if r_idx == 0: # Header
                            set_table_header_background(cell, '1B365D') # Deep Blue
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                            cell.paragraphs[0].runs[0].font.bold = True
                        elif r_idx % 2 == 0: # Zebra hatching
                            add_hatched_shading(cell, 'F2F2F2')
                            
                doc.add_paragraph() # Spacer
                table_data = []
                in_table = False
            continue
        
        # Horizontal Rule (fechar tabela pendente antes da quebra de página)
        if line == '---':
            if in_table and table_data:
                rows = len(table_data)
                cols = len(table_data[0])
                table = doc.add_table(rows=rows, cols=cols)
                table.style = "Table Grid"
                for r_idx, row_data in enumerate(table_data):
                    row = table.rows[r_idx]
                    for c_idx, text in enumerate(row_data):
                        cell = row.cells[c_idx]
                        cell.text = text
                        if r_idx == 0:
                            set_table_header_background(cell, "1B365D")
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                            cell.paragraphs[0].runs[0].font.bold = True
                        elif r_idx % 2 == 0:
                            add_hatched_shading(cell, "F2F2F2")
                doc.add_paragraph()
                table_data = []
                in_table = False
            doc.add_page_break()
            continue

        # Body Text
        if line and not in_table:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(12)

    # Última tabela do ficheiro (sem linha em branco a seguir)
    if in_table and table_data:
        rows = len(table_data)
        cols = len(table_data[0])
        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        for r_idx, row_data in enumerate(table_data):
            row = table.rows[r_idx]
            for c_idx, text in enumerate(row_data):
                cell = row.cells[c_idx]
                cell.text = text
                if r_idx == 0:
                    set_table_header_background(cell, "1B365D")
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                    cell.paragraphs[0].runs[0].font.bold = True
                elif r_idx % 2 == 0:
                    add_hatched_shading(cell, "F2F2F2")
        doc.add_paragraph()

    doc.save(docx_path)
    print(f"Livro gerado com sucesso em: {docx_path}")

if __name__ == "__main__":
    repo_root = Path(__file__).resolve().parents[1]
    # Fonte Markdown: canónico em 01_BOOKS; override com FLUXUS_ANNALS_MD se necessário
    MD_SOURCE = Path(
        os.environ.get(
            "FLUXUS_ANNALS_MD",
            str(repo_root / "01_BOOKS" / "BOOK_STRATEGIC_ANNALS_2026.md"),
        )
    )
    DOCX_OUTPUT = repo_root / "01_BOOKS" / "BOOK_STRATEGIC_ANNALS_AUTOGEN.docx"
    COVER_IMG = repo_root / "06_DATA" / "DATA_FLUXUS_COVER.png"

    generate_book(str(MD_SOURCE), str(DOCX_OUTPUT), str(COVER_IMG))

    # Réplicas com nomes usados na pauta comercial / sponsors
    for name in (
        "BOOK_STRATEGIC_ANNALS_2026.docx",
        "FLUXUS_STRATEGIC_ANNALS_2026.docx",
        "LIVRO_ESTRATEGICO.docx",
    ):
        dest = repo_root / "01_BOOKS" / name
        shutil.copyfile(DOCX_OUTPUT, dest)
        print(f"Copiado para: {dest}")
