"""Atualiza .docx FLUXUS: spread canónico 3,5% e repartição 42,8/28,6/28,6."""
from __future__ import annotations

from pathlib import Path

from docx import Document

REPO = Path(__file__).resolve().parents[1]

ANGEL_PARA_NEW = (
    "1.  **Spread P2P (Matching):** A diferença entre a taxa capturada do remetente e a taxa paga ao "
    'investidor de liquidez. Com o **spread total de referência de 3,5%** (documento canónico '
    "`03_TECHNICAL/TECH_FLUXUS_REMITTANCE_CANON.md`), o investidor FLUXUS opera em regime de "
    '"Lucro Presumido" e trilhos síncronos, competindo com Wise/Remitly e preservando ~**1,0%** ao '
    "protocolo (orquestração) **dentro** desse pacote, além de receitas de cartão onde aplicável."
)


def fix_book_angel_success_book_latest(path: Path) -> bool:
    d = Document(path)
    changed = False
    for p in d.paragraphs:
        if "margem de 1.5% a 2.5% para o protocolo" in p.text or (
            "Spread P2P (Matching)" in p.text and "2.5%" in p.text
        ):
            p.text = ANGEL_PARA_NEW
            changed = True
        if "**5% a 10%**" in p.text:
            p.text = p.text.replace(
                "O investidor ganha entre **5% a 10%** por operação ao prover a liquidez de última milha.",
                "O investidor captura a fatia **LP (~1,5% sobre o notional)** dentro do pacote de spread total "
                "de **3,5%** ao remetente (ver `03_TECHNICAL/TECH_FLUXUS_REMITTANCE_CANON.md`); o retorno anualizado "
                "(APY) depende do **número de ciclos** de liquidez, não desse valor isolado por operação.",
            )
            changed = True
        if p.text.startswith("Imagine que capturamos apenas") and "conservador):" in p.text:
            p.text = (
                "Imagine que capturamos apenas **0,005%** do mercado de remessas da Índia e Paquistão no primeiro "
                "ano (um objetivo extremamente conservador), com **spread total canónico de 3,5%** ao remetente:"
            )
            changed = True
        if p.text.startswith("- **Volume Transacional (GTV):** ~US$ 7,5"):
            p.text = "- **Volume Transacional (GTV):** ~US$ 15 Milhões/mês."
            changed = True
        if "Spread Gerado:** ~7%" in p.text:
            p.text = "- **Spread agregado (3,5% sobre GTV):** ~US$ 525.000,00/mês."
            changed = True
        if p.text.startswith("- **Taxa da Plataforma:** 1,5%"):
            p.text = "- **Orquestração (~1,0% sobre GTV):** ~US$ 150.000,00/mês."
            changed = True
    if changed:
        d.save(path)
    return changed


def fix_book_angel_remit_book(path: Path) -> bool:
    """Variante sem markdown pesado no cap. 3; cap. 8 com ** nos bullets."""
    d = Document(path)
    changed = False
    for p in d.paragraphs:
        if "O investidor ganha entre 5% a 10% por operação ao prover a liquidez de última milha." in p.text:
            p.text = p.text.replace(
                "O investidor ganha entre 5% a 10% por operação ao prover a liquidez de última milha.",
                "O investidor captura a fatia LP (~1,5% sobre o notional) dentro do pacote de spread total "
                "de 3,5% ao remetente (ver 03_TECHNICAL/TECH_FLUXUS_REMITTANCE_CANON.md); o retorno anualizado "
                "(APY) depende do número de ciclos de liquidez, não desse valor isolado por operação.",
            )
            changed = True
        if p.text.startswith("Imagine que capturamos apenas 0,005%") and p.text.endswith("conservador):"):
            p.text = (
                "Imagine que capturamos apenas 0,005% do mercado de remessas da Índia e Paquistão no primeiro "
                "ano (um objetivo extremamente conservador), com spread total canónico de 3,5% ao remetente:"
            )
            changed = True
        if p.text == "**Volume Transacional (GTV):** ~US$ 7,5 Milhões/mês.":
            p.text = "**Volume Transacional (GTV):** ~US$ 15 Milhões/mês."
            changed = True
        if p.text == "**Spread Gerado:** ~7% (US$ 525.000,00).":
            p.text = "**Spread agregado (3,5% sobre GTV):** ~US$ 525.000,00/mês."
            changed = True
        if p.text == "**Taxa da Plataforma:** 1,5% (US$ 112.500,00/mês).":
            p.text = "**Orquestração (~1,0% sobre GTV):** ~US$ 150.000,00/mês."
            changed = True
    if changed:
        d.save(path)
    return changed


def fix_strategic_annals(path: Path) -> bool:
    d = Document(path)
    changed = False
    mapping = [
        ("Para cada $100 de spread gerado:", "Para cada **$100** de spread total capturado (repartição canónica):"),
        ("- **$70** vão para o Investidor (LP).", "- **$42,80** para o Investidor (LP) (~42,8% do spread)."),
        ("- **$20** ficam com a Plataforma FLUXUS.", "- **$28,60** para o trilho VASP / local (~28,6% do spread)."),
        ("- **$10** vão para o Fundo de Reserva.", "- **$28,60** para a Plataforma FLUXUS (~28,6% do spread, orquestração)."),
    ]
    for p in d.paragraphs:
        for old, new in mapping:
            if old in p.text and new not in p.text:
                p.text = p.text.replace(old, new)
                changed = True

    if len(d.tables) > 1:
        row = d.tables[1].rows[1]
        if row.cells[0].text.strip() == "Brasil (PIX)" and "17,50" in row.cells[4].text:
            row.cells[3].text = "$ 7,01"
            row.cells[4].text = "R$ 15,00"
            changed = True

    if changed:
        d.save(path)
    return changed


def main() -> None:
    targets = [
        REPO / "01_BOOKS" / "BOOK_ANGEL_SUCCESS_BOOK_LATEST.docx",
        REPO / "01_BOOKS" / "BOOK_ANGEL_REMIT_BOOK.docx",
        REPO / "01_BOOKS" / "BOOK_STRATEGIC_ANNALS_2026.docx",
        REPO / "01_BOOKS" / "FLUXUS_STRATEGIC_ANNALS_2026.docx",
        REPO / "01_BOOKS" / "LIVRO_ESTRATEGICO.docx",
    ]
    for path in targets:
        if not path.exists():
            print("missing", path)
            continue
        if path.name == "BOOK_ANGEL_SUCCESS_BOOK_LATEST.docx":
            ok = fix_book_angel_success_book_latest(path)
        elif path.name == "BOOK_ANGEL_REMIT_BOOK.docx":
            ok = fix_book_angel_remit_book(path)
        else:
            ok = fix_strategic_annals(path)
        print(path.relative_to(REPO), "updated" if ok else "no change")


if __name__ == "__main__":
    main()
