import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_table_header_background(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def generate_mou_word(md_path, docx_path):
    print("Gerando MoU em formato Word...")
    doc = Document()
    
    # Text Styles
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(docx_path)
    print(f"MoU salvo em: {docx_path}")

if __name__ == "__main__":
    repo_root = Path(__file__).resolve().parents[1]
    MD_FILE = repo_root / "02_INVESTOR" / "INVESTOR_MOU_VASP_PARTNERSHIP.md"
    DOCX_FILE = repo_root / "02_INVESTOR" / "INVESTOR_MOU_VASP_PARTNERSHIP_AUTOGEN.docx"
    generate_mou_word(str(MD_FILE), str(DOCX_FILE))
