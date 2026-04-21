import os
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def merge_to_docx(book_dir, output_path, cover_image):
    doc = Document()
    
    # Customizing basic style
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # 1. Add Cover Image
    if os.path.exists(cover_image):
        doc.add_picture(cover_image, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

    # 2. Add Chapters (BOOK_ENCYCLOPEDIA_CAP*.md em 01_BOOKS; legado: FLUXUS_ENCYCLOPEDIA_CAP*.md)
    legacy = [f"FLUXUS_ENCYCLOPEDIA_CAP{i}.md" for i in range(1, 8)]
    found = sorted(Path(book_dir).glob("BOOK_ENCYCLOPEDIA_CAP*.md"), key=lambda p: p.name)
    if not found:
        found = [Path(book_dir) / name for name in legacy if (Path(book_dir) / name).exists()]

    for md_path in found:
        md_path = Path(md_path)
        if not md_path.exists():
            continue

        with open(md_path, "r", encoding="utf-8") as f:
            lines = f.readlines()

        for line in lines:
            line = line.strip()
            if not line or line.startswith('![Capa'):
                continue
                
            # Headers
            if line.startswith('### '):
                doc.add_heading(line.replace('###', '').strip(), level=3)
            elif line.startswith('## '):
                doc.add_heading(line.replace('##', '').strip(), level=2)
            elif line.startswith('# '):
                doc.add_heading(line.replace('#', '').strip(), level=1)
            
            # Bullet points
            elif line.startswith('- '):
                doc.add_paragraph(line.replace('-', '').strip(), style='List Bullet')
            
            # Horizontal rule / Page break
            elif line.startswith('---'):
                doc.add_page_break()
                
            else:
                p = doc.add_paragraph()
                # Handle bold text within lines
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part.replace('**', ''))
                        run.bold = True
                    else:
                        p.add_run(part)

    doc.save(output_path)
    print(f"Success: {output_path} created.")

if __name__ == "__main__":
    repo_root = Path(__file__).resolve().parents[1]
    book_folder = repo_root / "01_BOOKS"
    out_file = book_folder / "BOOK_FLUXUS_ENCYCLOPEDIA_MERGED.docx"
    cover = repo_root / "06_DATA" / "DATA_FLUXUS_COVER.png"
    merge_to_docx(str(book_folder), str(out_file), str(cover))
