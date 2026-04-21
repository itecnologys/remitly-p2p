import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_styled_whitepaper(md_file_path, output_path):
    doc = Document()
    
    # Custom Styles inspired by "eFin" Dribbble Pitch Decks (Sleek, Modern, High Contrast)
    styles = doc.styles
    
    # Title Style
    title_style = styles['Title']
    title_font = title_style.font
    title_font.name = 'Segoe UI Semibold'
    title_font.size = Pt(32)
    title_font.bold = True
    title_font.color.rgb = RGBColor(15, 23, 42) # Slate 900
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(15)

    # Heading 1 Style
    h1_style = styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Segoe UI Semibold'
    h1_font.size = Pt(22)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(30, 64, 175) # Blue 800
    h1_style.paragraph_format.space_before = Pt(24)
    h1_style.paragraph_format.space_after = Pt(12)
    # Add a border-bottom equivalent logic (optional, via xml) not required for simple elegance

    # Heading 2 Style
    h2_style = styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Segoe UI Semibold'
    h2_font.size = Pt(18)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(5, 150, 105) # Emerald 600
    h2_style.paragraph_format.space_before = Pt(18)
    h2_style.paragraph_format.space_after = Pt(8)

    # Heading 3 Style
    h3_style = styles['Heading 3']
    h3_font = h3_style.font
    h3_font.name = 'Segoe UI'
    h3_font.size = Pt(14)
    h3_font.bold = True
    h3_font.color.rgb = RGBColor(71, 85, 105) # Slate 600
    h3_style.paragraph_format.space_before = Pt(14)
    h3_style.paragraph_format.space_after = Pt(6)

    # Normal Text Style
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(51, 65, 85) # Slate 700
    normal_style.paragraph_format.line_spacing = 1.6
    normal_style.paragraph_format.space_after = Pt(10)
    
    # Add Cover Page
    doc.add_paragraph('\n\n')
    title = doc.add_paragraph('NEXUS PLATFORM', style='Title')
    subtitle = doc.add_paragraph('The Stablecoin P2P Lending Engine\nWhitepaper Executivo — V 1.0', style='Heading 2')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n\n\n\n\n\n\n\n')
    meta = doc.add_paragraph('Documentation structured after Institutional Pitch Templates', style='Normal')
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Read Markdown content
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line or line == '---':
            continue
            
        if line.startswith('# '):
            if 'NEXUS PLATFORM' not in line:
                doc.add_paragraph(line[2:].strip(), style='Heading 1')
        elif line.startswith('## '):
            doc.add_paragraph(line[3:].strip(), style='Heading 1')
        elif line.startswith('### '):
            p = doc.add_paragraph(style='Heading 2')
            # Extract icons if present: e.g. "📱 Dashboard"
            text = line[4:].strip()
            # simple bold replacement
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            text = line[2:]
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    # also replace backticks simply
                    part = part.replace('`', '')
                    p.add_run(part)
        elif line[0].isdigit() and len(line) > 1 and line[1] == '.':
            p = doc.add_paragraph(style='List Number')
            text = line[2:].strip()
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        else:
            p = doc.add_paragraph(style='Normal')
            text = line.replace('`', '')
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
                    
    doc.save(output_path)
    print(f"Whitepaper Document gerado com sucesso: {output_path}")

md_path = r'C:\LAPTOP\p2p-lending\p2p-lending by Antigravity\07_NEXUS_Stablecoin_Whitepaper.md'
docx_path = r'C:\LAPTOP\p2p-lending\p2p-lending by Antigravity\NEXUS_Whitepaper_Official.docx'

create_styled_whitepaper(md_path, docx_path)
