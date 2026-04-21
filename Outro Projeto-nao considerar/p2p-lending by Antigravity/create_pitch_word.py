import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def create_styled_document(md_file_path, output_path):
    doc = Document()
    
    # 1. Custom Styles inspired by Dribbble Pitch Decks (Sleek, Modern, Dark Blue/Emerald)
    styles = doc.styles
    
    # Title Style
    title_style = styles['Title']
    title_font = title_style.font
    title_font.name = 'Segoe UI'
    title_font.size = Pt(28)
    title_font.bold = True
    title_font.color.rgb = RGBColor(15, 23, 42) # Slate 900
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(20)

    # Heading 1 Style
    h1_style = styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Segoe UI'
    h1_font.size = Pt(20)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(30, 58, 138) # Blue 900
    h1_style.paragraph_format.space_before = Pt(24)
    h1_style.paragraph_format.space_after = Pt(12)

    # Heading 2 Style
    h2_style = styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Segoe UI'
    h2_font.size = Pt(16)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(4, 120, 87) # Emerald 700
    h2_style.paragraph_format.space_before = Pt(18)
    h2_style.paragraph_format.space_after = Pt(8)

    # Normal Text Style
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(51, 65, 85) # Slate 700
    normal_style.paragraph_format.line_spacing = 1.5

    # Quote/Insight Style
    try:
        quote_style = styles.add_style('QuoteInsight', 1) # 1 = PARAGRAPH
    except:
        quote_style = styles['QuoteInsight']
    quote_font = quote_style.font
    quote_font.name = 'Georgia'
    quote_font.size = Pt(12)
    quote_font.italic = True
    quote_font.color.rgb = RGBColor(217, 119, 6) # Amber 600
    quote_style.paragraph_format.left_indent = Inches(0.5)
    quote_style.paragraph_format.right_indent = Inches(0.5)
    quote_style.paragraph_format.space_before = Pt(10)
    quote_style.paragraph_format.space_after = Pt(10)
    
    # Add Cover Page
    doc.add_paragraph('\n')
    doc.add_paragraph('\n')
    title = doc.add_paragraph('NEXUS PLATFORM', style='Title')
    subtitle = doc.add_paragraph('Thesis de Investimento & Manual de Controladoria V8', style='Heading 1')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n\n\n')
    meta = doc.add_paragraph('Uso Exclusivo: Board Executivo e C-Level\nConfidencial', style='Normal')
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Read Markdown content
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    in_quote = False
    
    for line in lines:
        line = line.strip()
        if not line or line == '---':
            continue
            
        if line.startswith('# '):
            # Skip the markdown title since we have a cover
            if 'MANUAL' not in line:
                doc.add_paragraph(line[2:].strip(), style='Heading 1')
        elif line.startswith('## '):
            doc.add_paragraph(line[3:].strip(), style='Heading 1')
        elif line.startswith('### '):
            doc.add_paragraph(line[4:].strip(), style='Heading 2')
        elif line.startswith('>'):
            in_quote = True
            doc.add_paragraph(line[1:].strip().replace('*', ''), style='QuoteInsight')
        elif line.startswith('* '):
            # Bullet point
            p = doc.add_paragraph(style='List Bullet')
            text = line[2:]
            # Simple bold parsing: **text**
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        else:
            in_quote = False
            p = doc.add_paragraph(style='Normal')
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
                    
    doc.save(output_path)
    print(f"Documento gerado com sucesso: {output_path}")

md_path = r'd:\LAPTOP\p2p-lending\p2p-lending by Antigravity\06_Manual_NEXUS_Controladoria_e_Tese.md'
docx_path = r'd:\LAPTOP\p2p-lending\p2p-lending by Antigravity\NEXUS_Pitch_Controladoria_V8.docx'

create_styled_document(md_path, docx_path)
