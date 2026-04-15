import os
from docx import Document
from docx.shared import Pt
import re

def markdown_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    
    # Customizing basic style
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        
        if not line:
            continue
            
        # Headers
        if line.startswith('### '):
            doc.add_heading(line.replace('###', '').strip(), level=3)
        elif line.startswith('## '):
            doc.add_heading(line.replace('##', '').strip(), level=2)
        elif line.startswith('# '):
            doc.add_heading(line.replace('#', '').strip(), level=1)
        
        # Bullet points
        elif line.startswith('- '):
            doc.add_paragraph(line.replace('-', '').strip(), style='List Bullet')
        
        # Bold text (simple replacement for the whole line if bolded)
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line.replace('**', '').strip())
            run.bold = True
        
        # Horizontal rule
        elif line.startswith('---'):
            doc.add_page_break()
            
        else:
            # Simple bold search in the middle of line
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part.replace('**', ''))
                    run.bold = True
                else:
                    p.add_run(part)

    doc.save(docx_path)
    print(f"Success: {docx_path} created.")

if __name__ == "__main__":
    md_file = r"c:\LAPTOP\remitly-p2p\p2p-remitly by Antigravity\FLUXUS_ANGEL_SUCCESS_BOOK.md"
    docx_file = r"c:\LAPTOP\remitly-p2p\p2p-remitly by Antigravity\FLUXUS_ANGEL_REMIT_BOOK.docx"
    markdown_to_docx(md_file, docx_file)
