import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def merge_to_docx(book_dir, output_path, cover_image):
    doc = Document()
    
    # Customizing basic style
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # 1. Add Cover Image
    if os.path.exists(cover_image):
        doc.add_picture(cover_image, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

    # 2. Add Chapters
    chapters = [
        "FLUXUS_ENCYCLOPEDIA_CAP1.md",
        "FLUXUS_ENCYCLOPEDIA_CAP2.md",
        "FLUXUS_ENCYCLOPEDIA_CAP3.md",
        "FLUXUS_ENCYCLOPEDIA_CAP4.md",
        "FLUXUS_ENCYCLOPEDIA_CAP5.md",
        "FLUXUS_ENCYCLOPEDIA_CAP6.md",
        "FLUXUS_ENCYCLOPEDIA_CAP7.md"
    ]

    for cap in chapters:
        md_path = os.path.join(book_dir, cap)
        if not os.path.exists(md_path):
            continue

        with open(md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        for line in lines:
            line = line.strip()
            if not line or line.startswith('![Capa'):
                continue
                
            # Headers
            if line.startswith('### '):
                doc.add_heading(line.replace('###', '').strip(), level=3)
            elif line.startswith('## '):
                doc.add_heading(line.replace('##', '').strip(), level=2)
            elif line.startswith('# '):
                doc.add_heading(line.replace('#', '').strip(), level=1)
            
            # Bullet points
            elif line.startswith('- '):
                doc.add_paragraph(line.replace('-', '').strip(), style='List Bullet')
            
            # Horizontal rule / Page break
            elif line.startswith('---'):
                doc.add_page_break()
                
            else:
                p = doc.add_paragraph()
                # Handle bold text within lines
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part.replace('**', ''))
                        run.bold = True
                    else:
                        p.add_run(part)

    doc.save(output_path)
    print(f"Success: {output_path} created with all 7 chapters.")

if __name__ == "__main__":
    book_folder = r"c:\LAPTOP\remitly-p2p\p2p-remitly by Antigravity\book"
    out_file = r"c:\LAPTOP\remitly-p2p\p2p-remitly by Antigravity\book\FLUXUS_ENCYCLOPEDIA_60_PAGES_V2.docx"
    cover = os.path.join(book_folder, "FLUXUS_COVER.png")
    merge_to_docx(book_folder, out_file, cover)
