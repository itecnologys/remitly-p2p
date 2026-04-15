import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')

    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def convert_md_to_docx(md_path, docx_path, cover_path):
    doc = Document()

    # Define standard styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Add Cover Image
    if os.path.exists(cover_path):
        doc.add_picture(cover_path, width=Inches(6.5))
        doc.add_page_break()

    # Add Header
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.text = "FLUXUS: THE ENGINE — Manual Técnico de Orquestração"
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add Footer with Page Number
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    run.add_text("Página ")
    add_page_number(run)
    run.add_text(" de ")
    # For total pages we use NUMPAGES but let's keep it simple with just PAGE for now
    # as NUMPAGES requires more complex XML
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Split by chapters/pages
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            # Check if it's a "Página X" header to add page break
            if "Página" in line:
                doc.add_page_break()
            doc.add_heading(line[4:], level=2)
        elif line == '---':
            # Horizontal rule could be a page break or a separator
            # In this book context, let's treat it as a light separator
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("_________________________________________________")
        else:
            # Body text
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5

    doc.save(docx_path)
    print(f"Documento salvo em: {docx_path}")

if __name__ == "__main__":
    MD_FILE = r'c:\LAPTOP\remitly-p2p\p2p-remitly by Antigravity\engine\FLUXUS_ENGINE_BOOK.md'
    DOCX_FILE = r'c:\LAPTOP\remitly-p2p\p2p-remitly by Antigravity\engine\FLUXUS_ENGINE_BOOK.docx'
    COVER_IMAGE = r'C:\Users\mini\.gemini\antigravity\brain\ac6c47b0-eeff-4199-a4d8-f32eb965751f\fluxus_engine_gold_cover_1776194239234.png'
    
    convert_md_to_docx(MD_FILE, DOCX_FILE, COVER_IMAGE)
